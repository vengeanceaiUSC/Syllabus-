#!/usr/bin/env python3
"""Extract plain text from a syllabus file.

Supports PDF (.pdf), Word (.docx), HTML (.html/.htm), and plain text
(.txt/.md). Prints the extracted text to stdout so the agent can read and
structure it.

Usage:
    python extract_text.py <path-to-syllabus>
    python extract_text.py <path-to-syllabus> --out extracted.txt
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

CALENDAR_TABLE_START = "=== CALENDAR_TABLE_JSON ==="
CALENDAR_TABLE_END = "=== END CALENDAR_TABLE ==="


def _clean_cell(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"\s+", " ", value.replace("\n", " / ").strip())


def _parse_calendar_tables(pdf) -> list[dict]:
    """Extract structured rows from PDF table layouts (Course Calendar format)."""
    rows: list[dict] = []
    seen: set[tuple[str, ...]] = set()

    for page in pdf.pages:
        for table in page.extract_tables() or []:
            if len(table) < 2:
                continue
            header = " ".join(_clean_cell(c) for c in table[0]).lower()
            if "homework" not in header and "topic" not in header:
                continue
            for raw in table[1:]:
                cells = list(raw) + [None] * max(0, 7 - len(raw))
                row = {
                    "class": _clean_cell(cells[0]),
                    "date_mon": _clean_cell(cells[1]),
                    "date_wed": _clean_cell(cells[2]),
                    "topic": _clean_cell(cells[3]),
                    "reading": _clean_cell(cells[5]),
                    "homework": _clean_cell(cells[6]),
                }
                key = tuple(row.values())
                if not any(row.values()) or key in seen:
                    continue
                seen.add(key)
                rows.append(row)
    return rows


def extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(
            "pdfplumber is required for PDF files. Install with: pip install pdfplumber"
        ) from exc

    pages: list[str] = []
    calendar_rows: list[dict] = []
    with pdfplumber.open(str(path)) as pdf:
        calendar_rows = _parse_calendar_tables(pdf)
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            pages.append(text)
            pages.append(f"\n-- {i} of {len(pdf.pages)} --\n")
    body = "\n".join(pages)
    if calendar_rows:
        body += (
            f"\n{CALENDAR_TABLE_START}\n"
            f"{json.dumps(calendar_rows, ensure_ascii=False)}\n"
            f"{CALENDAR_TABLE_END}\n"
        )
    return body


def extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(
            "python-docx is required for .docx files. Install with: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_html(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        # Fallback: crude tag stripping.
        import re

        return re.sub(r"<[^>]+>", " ", raw)

    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return soup.get_text(separator="\n")


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".html", ".htm"}:
        return extract_html(path)
    if suffix in {".txt", ".md", ".text"}:
        return path.read_text(encoding="utf-8", errors="replace")
    raise SystemExit(
        f"Unsupported file type '{suffix}'. Supported: .pdf, .docx, .html, .htm, .txt, .md"
    )


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract text from a syllabus file.")
    parser.add_argument("path", help="Path to the syllabus file")
    parser.add_argument("--out", help="Optional path to write the extracted text")
    args = parser.parse_args()

    path = Path(args.path).expanduser()
    if not path.exists():
        raise SystemExit(f"File not found: {path}")

    text = extract_text(path)

    if args.out:
        Path(args.out).expanduser().write_text(text, encoding="utf-8")
        print(f"Wrote extracted text to {args.out} ({len(text)} chars)")
    else:
        sys.stdout.write(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
